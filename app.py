# ============================================================
# ValidAI Risk — M4
# Copiloto IA para validación de modelos de riesgo
# Pipeline completo:
#   KB -> Loader -> Limpieza -> Chunker -> Metadata ->
#   Embeddings -> ES Hybrid (BM25+kNN+RRF) -> Normativa SBS ->
#   HyDE Retriever -> GPT-4o-mini -> Reporte -> HITL -> SQLite
# ============================================================
# ============================================================
# 0. Importaciones
# ============================================================

import os
import tempfile
import re
import io
import json
import uuid
import time
import zipfile
import sqlite3
import traceback

from datetime import datetime
from pathlib import Path
from typing import Optional

import streamlit as st
import pandas as pd
import numpy as np
from sklearn.metrics import roc_auc_score
from pypdf import PdfReader
from docx import Document
from pydantic import BaseModel, Field
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.messages import SystemMessage, HumanMessage

try:
    from langchain.agents import create_agent
    AGENT_BACKEND = "langchain_create_agent"
except Exception:
    create_agent = None
    AGENT_BACKEND = "langgraph_create_react_agent"
    from langgraph.prebuilt import create_react_agent

from langchain_core.documents import Document as LCDocument
from langchain_community.document_loaders import (
    PyPDFLoader, Docx2txtLoader, TextLoader, CSVLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from elasticsearch import Elasticsearch

# ============================================================
# 1. CONFIGURACIÓN GENERAL
# ============================================================

APP_NAME        = "ValidAI Risk"
DB_ENGINE       = "sqlite"
DB_NAME         = "memoria_validairisk.db"
MAX_CONTEXT_CHARS = 10000
INDEX_NAME      = "validairisk_hybrid"

FORMATOS_METODOLOGIA = ["pdf", "docx", "txt", "md"]
FORMATOS_CODIGO      = ["py", "ipynb", "sql", "txt", "md", "json", "zip"]
FORMATOS_DATOS       = ["csv", "xlsx", "xls"]
EXTENSIONES_CODIGO_ZIP = [".py", ".ipynb", ".sql", ".txt", ".md", ".json"]

DOMINIO_PERMITIDO = [
    "validacion", "validacion", "modelo", "riesgo", "scoring", "score",
    "banca", "metodologia", "metodologia", "codigo", "codigo", "python",
    "poblacion objetivo", "poblacion objetivo", "metricas", "metricas",
    "gini", "auc", "psi", "ks", "benchmark", "hallazgo", "observacion",
    "observacion", "recomendacion", "recomendacion", "variables",
    "calibracion", "calibracion", "performance", "monitoreo", "athena",
    "sql", "target", "filtro", "mora", "drift", "estabilidad",
    "reentrenamiento", "incumplimiento", "probabilidad de incumplimiento",
    "pd", "risk", "sbs", "normativa", "resolucion", "resolucion",
    "articulo", "articulo", "validacion independiente"
]

INDEX_MAPPING = {
    "settings": {
        "number_of_shards": 1, "number_of_replicas": 0,
        "analysis": {
            "analyzer": {"spanish_custom": {
                "type": "custom", "tokenizer": "standard",
                "filter": ["lowercase", "spanish_stop"]
            }},
            "filter": {"spanish_stop": {"type": "stop", "stopwords": "_spanish_"}}
        }
    },
    "mappings": {
        "properties": {
            "page_content": {"type": "text", "analyzer": "spanish_custom"},
            "embedding":    {"type": "dense_vector", "dims": 1536,
                             "index": True, "similarity": "cosine"},
            "metadata": {"properties": {
                "source":        {"type": "keyword"},
                "page":          {"type": "integer"},
                "tipo_documento":{"type": "keyword"},
                "caso_id":       {"type": "keyword"},
                "normativa_sbs": {"type": "keyword"},
                "articulo_sbs":  {"type": "keyword"},
                "fecha":         {"type": "date", "format": "yyyy-MM-dd"},
                "version":       {"type": "keyword"},
                "topics":        {"type": "keyword"}
            }}
        }
    }
}

# 2. CONFIGURACIÓN DE PÁGINA
# ============================================================

st.set_page_config(
    page_title=APP_NAME,
    layout="wide"
)


# ============================================================

# 3. ESTILO VISUAL
# ============================================================

def aplicar_estilo_visual():
    st.markdown("""
    <style>

    .stApp {
        background: #050505;
        color: #F5F5F5;
        font-family: "Inter", "IBM Plex Sans", sans-serif;
    }

    section[data-testid="stSidebar"] {
        background: #0A0A0A;
        border-right: 1px solid #262626;
    }

    h1, h2, h3 {
        color: #FFFFFF !important;
        font-weight: 850 !important;
        letter-spacing: -0.03em;
    }

    h4 {
        color: #D4D4D4 !important;
        font-weight: 750 !important;
    }

    p, label, span, div {
        color: #E5E5E5;
    }

    hr {
        border-color: #262626 !important;
    }

    button[data-baseweb="tab"] {
        background-color: #0A0A0A !important;
        border: 1px solid #262626 !important;
        border-radius: 14px !important;
        color: #A3A3A3 !important;
        margin-right: 6px !important;
        font-weight: 750 !important;
        padding: 9px 16px !important;
    }

    button[data-baseweb="tab"][aria-selected="true"] {
        background-color: #171717 !important;
        color: #FFFFFF !important;
        border: 1px solid #22C55E !important;
    }

    .stButton > button, .stDownloadButton > button {
        background-color: #111111 !important;
        color: #FFFFFF !important;
        border: 1px solid #404040 !important;
        border-radius: 14px !important;
        padding: 0.7rem 1.05rem !important;
        font-weight: 800 !important;
        transition: all 0.15s ease-in-out;
    }

    .stButton > button:hover, .stDownloadButton > button:hover {
        background-color: #1A1A1A !important;
        border: 1px solid #22C55E !important;
        transform: translateY(-1px);
    }

    .stTextInput input, .stTextArea textarea {
        background-color: #111111 !important;
        color: #FFFFFF !important;
        border: 1px solid #404040 !important;
        border-radius: 14px !important;
    }

    div[data-testid="stFileUploader"] {
        background-color: #111111;
        border: 1px dashed #404040;
        border-radius: 18px;
        padding: 12px;
    }

    div[data-testid="stFileUploader"] small {
        display: none !important;
    }

    div[data-testid="stFileUploader"] section div div span {
        display: none !important;
    }

    div[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzoneInstructions"] > div:nth-child(2) {
        display: none !important;
    }

    div[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzoneInstructions"] small {
        display: none !important;
    }

    details {
        background-color: #111111 !important;
        border: 1px solid #262626 !important;
        border-radius: 16px !important;
    }

    pre {
        background-color: #050505 !important;
        border: 1px solid #262626 !important;
        border-radius: 14px !important;
        color: #E5E5E5 !important;
    }

    div[data-testid="stMetric"] {
        background-color: #111111;
        border: 1px solid #262626;
        border-radius: 16px;
        padding: 12px;
    }

    .risk-card {
        border: 1px solid #262626;
        border-radius: 20px;
        padding: 18px;
        background: #111111;
        min-height: 175px;
    }

    .risk-card-title {
        color: #FFFFFF;
        font-size: 1rem;
        font-weight: 900;
        margin-bottom: 8px;
    }

    .risk-card-text {
        color: #BDBDBD;
        font-size: 0.88rem;
        line-height: 1.45;
    }

    .soft-badge {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        background: #171717;
        color: #D4D4D4;
        border: 1px solid #404040;
        font-size: 0.76rem;
        font-weight: 800;
        margin-bottom: 8px;
    }

    .green-badge {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        background: rgba(34, 197, 94, 0.08);
        color: #A7F3D0;
        border: 1px solid rgba(34, 197, 94, 0.35);
        font-size: 0.76rem;
        font-weight: 800;
        margin-bottom: 8px;
    }

    .mini-note {
        border: 1px solid #262626;
        border-radius: 16px;
        padding: 12px;
        background: #111111;
        color: #BDBDBD;
        font-size: 0.9rem;
    }

    .status-ok {
        color: #A7F3D0 !important;
        font-weight: 800;
    }

    .status-pending {
        color: #D4D4D4 !important;
        font-weight: 800;
    }

    .status-error {
        color: #FCA5A5 !important;
        font-weight: 800;
    }

    </style>
    """, unsafe_allow_html=True)


def render_logo_header():
    """
    Header 100% seguro sin HTML en el contenido visible.
    """

    st.title("ValidAI Risk")

    st.caption(
        "Copiloto IA para validación metodológica, revisión de código, "
        "trazabilidad y memoria en modelos de riesgo."
    )


def card_html(icono: str, titulo: str, texto: str, badge: str = "MVP") -> str:
    return f"""
    <div class="risk-card">
        <div class="soft-badge">{badge}</div>
        <div class="risk-card-title">{icono} {titulo}</div>
        <div class="risk-card-text">{texto}</div>
    </div>
    """


# ============================================================

# 4. UTILIDADES
# ============================================================

def configurar_apis():

    credenciales = {}

    try:

        with open(

            "api.txt",

            "r",

            encoding="utf-8"

        ) as f:

            for linea in f:

                if "=" in linea:

                    clave, valor = linea.strip().split(

                        "=",

                        1

                    )

                    credenciales[

                        clave.strip()

                    ] = valor.strip()

    except Exception:

        return


    openai_key = credenciales.get(

        "OPENAI_API_KEY",

        ""

    )

    langsmith_key = credenciales.get(

        "LANGSMITH_API_KEY",

        ""

    )


    if openai_key:

        os.environ[

            "OPENAI_API_KEY"

        ] = openai_key


    if langsmith_key:

        os.environ[

            "LANGCHAIN_TRACING_V2"

        ] = "true"

        os.environ[

            "LANGCHAIN_API_KEY"

        ] = langsmith_key

        os.environ[

            "LANGCHAIN_PROJECT"

        ] = "M2-Copiloto-Validacion-Modelos"

        os.environ[

            "LANGCHAIN_ENDPOINT"

        ] = "https://api.smith.langchain.com"

    else:

        os.environ[

            "LANGCHAIN_TRACING_V2"

        ] = "false"


def limpiar_texto(texto: str) -> str:

    if not texto:

        return ""


    texto = texto.replace(

        "\x00",

        " "

    )


    texto = re.sub(

        r"\s+",

        " ",

        texto

    )


    return texto.strip()


def anonimizar_texto_basico(texto: str) -> str:

    if not texto:

        return ""


    texto = re.sub(

        r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",

        "[EMAIL]",

        texto

    )


    texto = re.sub(

        r"\b\d{8,12}\b",

        "[ID_NUMERICO]",

        texto

    )


    texto = re.sub(

        r"\b9\d{8}\b",

        "[TELEFONO]",

        texto

    )


    texto = re.sub(

        r"\b\d{13,19}\b",

        "[NUMERO_LARGO]",

        texto

    )


    return texto


def limitar_contexto(

    texto: str,

    max_chars: int = MAX_CONTEXT_CHARS

) -> str:


    if not texto:

        return ""


    texto = limpiar_texto(

        texto

    )


    if len(texto) <= max_chars:

        return texto


    return (

        texto[:max_chars]

        + "\n\n[CONTEXTO LIMITADO: el archivo supera el tamaño máximo configurado para el MVP. "

        + "En producción se recomienda usar chunking/RAG para analizar el documento completo.]"

    )


def preparar_contexto(

    texto: str,

    max_chars: int = MAX_CONTEXT_CHARS

) -> str:


    texto = limpiar_texto(

        texto

    )


    texto = anonimizar_texto_basico(

        texto

    )


    texto = limitar_contexto(

        texto,

        max_chars=max_chars

    )


    return texto


def leer_archivo(uploaded_file) -> str:
    """
    Lee un archivo subido via st.file_uploader (PDF, DOCX o TXT) y
    devuelve su contenido como texto plano. Usado para el documento
    metodologico del modelo.
    """
    if uploaded_file is None:
        return ""

    nombre = uploaded_file.name.lower()
    sufijo = "." + nombre.split(".")[-1] if "." in nombre else ""

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=sufijo) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        if sufijo == ".pdf":
            paginas = PyPDFLoader(tmp_path).load()
            texto = "\n\n".join(p.page_content for p in paginas)
        elif sufijo == ".docx":
            texto = Docx2txtLoader(tmp_path).load()[0].page_content
        elif sufijo in [".txt", ".md"]:
            texto = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        else:
            texto = uploaded_file.getvalue().decode("utf-8", errors="ignore")

        os.unlink(tmp_path)
        return texto

    except Exception as e:
        return f"[ERROR_LECTURA_ARCHIVO] {uploaded_file.name}: {str(e)}"


def leer_multiples_archivos_codigo(archivos_subidos) -> str:
    """
    Lee uno o varios archivos de código/notebook subidos via
    st.file_uploader (.py, .ipynb, .sql, .zip) y devuelve el
    contenido concatenado como texto plano, identificando cada
    archivo con un encabezado.
    """
    if not archivos_subidos:
        return ""

    bloques = []
    for archivo in archivos_subidos:
        nombre = archivo.name.lower()
        try:
            if nombre.endswith(".ipynb"):
                nb_json = json.loads(archivo.getvalue().decode("utf-8", errors="ignore"))
                celdas_texto = []
                for celda in nb_json.get("cells", []):
                    fuente = "".join(celda.get("source", []))
                    if celda.get("cell_type") == "code":
                        celdas_texto.append(f"# [CODE]\n{fuente}")
                    else:
                        celdas_texto.append(f"# [MARKDOWN]\n{fuente}")
                contenido = "\n\n".join(celdas_texto)
            elif nombre.endswith(".zip"):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
                    tmp.write(archivo.getvalue())
                    tmp_path = tmp.name
                contenido = leer_zip_codigo(tmp_path)
                os.unlink(tmp_path)
            else:
                contenido = archivo.getvalue().decode("utf-8", errors="ignore")

            bloques.append(f"===== ARCHIVO: {archivo.name} =====\n{contenido}")

        except Exception as e:
            bloques.append(f"===== ARCHIVO: {archivo.name} =====\n[ERROR_LECTURA] {str(e)}")

    return "\n\n".join(bloques)


def leer_zip_codigo(ruta_zip: str) -> str:
    """
    Extrae y concatena el contenido de archivos .py, .sql e .ipynb
    dentro de un .zip de código fuente del modelo.
    """
    bloques = []
    try:
        with zipfile.ZipFile(ruta_zip, "r") as z:
            for nombre_interno in z.namelist():
                if nombre_interno.endswith((".py", ".sql", ".ipynb")):
                    try:
                        contenido = z.read(nombre_interno).decode("utf-8", errors="ignore")
                        bloques.append(f"--- {nombre_interno} ---\n{contenido}")
                    except Exception:
                        continue
    except Exception as e:
        return f"[ERROR_LECTURA_ZIP] {str(e)}"
    return "\n\n".join(bloques)


def leer_excel_o_csv(uploaded_file):
    """
    Lee un archivo de métricas subido via st.file_uploader (.xlsx, .xls o .csv)
    y devuelve un DataFrame de pandas, o None si falla la lectura.
    """
    if uploaded_file is None:
        return None

    nombre = uploaded_file.name.lower()
    try:
        if nombre.endswith(".csv"):
            return pd.read_csv(uploaded_file)
        elif nombre.endswith((".xlsx", ".xls")):
            return pd.read_excel(uploaded_file)
        else:
            return None
    except Exception as e:
        st.error(f"Error al leer {uploaded_file.name}: {str(e)}")
        return None


def validar_dominio(

    texto: str

) -> bool:


    if not texto:

        return False


    t = texto.lower()


    return any(

        palabra in t

        for palabra in DOMINIO_PERMITIDO

    )


def fallback_falta_info(

    campo: str

) -> str:


    return (

        f"No se encontró información suficiente para analizar: {campo}. "

        f"Adjunta o completa ese insumo."

    )


def safe_json_loads(

    texto: str,

    default=None

):


    try:

        return json.loads(

            texto

        )


    except Exception:

        return default

# ============================================================


# ============================================================
# 4b. INICIALIZACION ELASTICSEARCH + EMBEDDINGS + RETRIEVERS
# ============================================================

@st.cache_resource
def init_elasticsearch():
    creds = configurar_apis_dict()
    url = creds.get("ELASTIC_URL", "")
    key = creds.get("ELASTIC_API_KEY", "")
    if url and key:
        try:
            client = Elasticsearch(url, api_key=key)
            client.info()
            return client
        except Exception:
            return None
    return None

@st.cache_resource
def init_embeddings_model():
    if os.environ.get("OPENAI_API_KEY"):
        try:
            return OpenAIEmbeddings(model="text-embedding-3-small")
        except Exception:
            return None
    return None

def configurar_apis_dict():
    creds = {}
    try:
        with open("api.txt", "r", encoding="utf-8") as f:
            for linea in f:
                if "=" in linea:
                    k, v = linea.strip().split("=", 1)
                    creds[k.strip()] = v.strip()
    except Exception:
        pass
    for key in ["OPENAI_API_KEY","ELASTIC_URL","ELASTIC_API_KEY","LANGSMITH_API_KEY"]:
        if key in creds:
            os.environ[key] = creds[key]
    if creds.get("LANGSMITH_API_KEY"):
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"]     = creds["LANGSMITH_API_KEY"]
        os.environ["LANGCHAIN_PROJECT"]     = "ValidAIRisk-M4-Hybrid"
        os.environ["LANGCHAIN_ENDPOINT"]    = "https://api.smith.langchain.com"
    else:
        os.environ["LANGCHAIN_TRACING_V2"] = "false"
    return creds

def check_rrf(es_client):
    if es_client is None:
        return False
    try:
        v = tuple(int(x) for x in es_client.info()["version"]["number"].split(".")[:2])
        return v >= (8, 9)
    except Exception:
        return False

def crear_indice_si_no_existe(es_client):
    if es_client is None:
        return
    try:
        if not es_client.indices.exists(index=INDEX_NAME):
            es_client.indices.create(index=INDEX_NAME, body=INDEX_MAPPING)
            log_debug("ES", f"Indice {INDEX_NAME} creado")
    except Exception as e:
        log_debug("ES", f"Error creando indice: {e}", "Error")

# 5. TRAZABILIDAD LOCAL
# ============================================================

def inicializar_debug():

    if "debug_logs" not in st.session_state:

        st.session_state["debug_logs"] = []


    if "ultimo_prompt" not in st.session_state:

        st.session_state["ultimo_prompt"] = ""


    if "ultimo_resumen_metricas" not in st.session_state:

        st.session_state["ultimo_resumen_metricas"] = ""


    if "ultimo_error" not in st.session_state:

        st.session_state["ultimo_error"] = ""


    if "estado_flujo" not in st.session_state:

        st.session_state["estado_flujo"] = {

            "insumos":"Pendiente",

            "preparacion":"Pendiente",

            "rag":"Pendiente",

            "evidencia":"Pendiente",

            "metricas":"Pendiente",

            "memoria":"Pendiente",

            "agente":"Pendiente",

            "reporte":"Pendiente"

        }


def log_debug(

    evento: str,

    detalle: str = "",

    estado: str = "OK"

):

    if "debug_logs" not in st.session_state:

        st.session_state["debug_logs"] = []


    st.session_state["debug_logs"].append(

        {

            "hora": datetime.now().strftime(

                "%H:%M:%S"

            ),

            "estado": estado,

            "evento": evento,

            "detalle": detalle

        }

    )


def actualizar_estado_flujo(

    paso: str,

    estado: str

):

    if "estado_flujo" not in st.session_state:

        inicializar_debug()


    st.session_state["estado_flujo"][

        paso

    ] = estado


def limpiar_debug():

    st.session_state["debug_logs"] = []


    st.session_state["ultimo_prompt"] = ""


    st.session_state["ultimo_resumen_metricas"] = ""


    st.session_state["ultimo_error"] = ""


    st.session_state["estado_flujo"] = {

        "insumos":"Pendiente",

        "preparacion":"Pendiente",

        "rag":"Pendiente",

        "evidencia":"Pendiente",

        "metricas":"Pendiente",

        "memoria":"Pendiente",

        "agente":"Pendiente",

        "reporte":"Pendiente"

    }


# ============================================================


# ============================================================
# 5b. RAGTRACER SQLite — Trazabilidad RAG
# ============================================================

def init_rag_traces_table():
    conn = sqlite3.connect(DB_NAME)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS rag_traces (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT, pregunta TEXT, hyde_doc TEXT,
            modo TEXT, n_chunks INTEGER, fuentes TEXT,
            normativas TEXT, latencia_ms INTEGER, error TEXT)
    """)
    conn.commit()
    conn.close()

def log_rag_trace(pregunta, hyde_doc, modo, resultados, latencia_ms, error=None):
    try:
        fuentes    = ", ".join({r["metadata"].get("source","?") for r in resultados})
        normativas = ", ".join({r["metadata"].get("normativa_sbs","")
                                for r in resultados if r["metadata"].get("normativa_sbs")})
        conn = sqlite3.connect(DB_NAME)
        conn.execute(
            """INSERT INTO rag_traces
               (timestamp,pregunta,hyde_doc,modo,n_chunks,fuentes,normativas,latencia_ms,error)
               VALUES (?,?,?,?,?,?,?,?,?)""",
            (datetime.now().isoformat(), pregunta[:500], (hyde_doc or "")[:500],
             modo, len(resultados), fuentes[:300], normativas[:200], latencia_ms, error)
        )
        conn.commit()
        conn.close()
    except Exception:
        pass

def get_rag_traces(limite=10):
    try:
        conn = sqlite3.connect(DB_NAME)
        df   = pd.read_sql_query(
            "SELECT * FROM rag_traces ORDER BY timestamp DESC LIMIT ?",
            conn, params=(limite,))
        conn.close()
        return df
    except Exception:
        return pd.DataFrame()

# 6. LECTURA DE ARCHIVOS
# ============================================================

class DocumentLoader:

    def __init__(self, knowledge_path):

        self.knowledge_path = Path(knowledge_path)


    def load_documents(self):

        documents = []

        for file in self.knowledge_path.rglob("*"):

            if not file.is_file():

                continue

            extension = file.suffix.lower()

            try:

                if extension == ".pdf":

                    loader = PyPDFLoader(str(file))

                    documents.extend(

                        loader.load()

                    )


                elif extension == ".docx":

                    loader = Docx2txtLoader(str(file))

                    documents.extend(

                        loader.load()

                    )


                elif extension in [

                    ".txt",

                    ".md",

                    ".py",

                    ".sql"

                ]:

                    loader = TextLoader(

                        str(file),

                        encoding="utf-8"

                    )

                    documents.extend(

                        loader.load()

                    )


                elif extension == ".csv":

                    loader = CSVLoader(

                        str(file)

                    )

                    documents.extend(

                        loader.load()

                    )


                elif extension in [

                    ".xlsx",

                    ".xls"

                ]:

                    documents.extend(

                        self.load_excel(

                            file

                        )

                    )


                elif extension == ".ipynb":

                    documents.extend(

                        self.load_notebook(

                            file

                        )

                    )


                elif extension == ".zip":

                    documents.extend(

                        self.load_zip(

                            file

                        )

                    )


            except Exception as e:

                print(

                    f"Error leyendo {file}: {e}"

                )

        return documents


    def load_excel(

        self,

        file

    ):

        df = pd.read_excel(file)

        text = df.to_string(

            index=False

        )

        return [

            Document(

                page_content=text,

                metadata={

                    "source":

                    file.name

                }

            )

        ]


    def load_notebook(

        self,

        file

    ):

        with open(

            file,

            encoding="utf-8"

        ) as f:

            notebook = json.load(f)


        docs = []


        for idx, cell in enumerate(

            notebook.get("cells",[])

        ):


            if cell["cell_type"] == "code":


                content = "".join(

                    cell["source"]

                )


                docs.append(

                    Document(

                        page_content=content,

                        metadata={

                            "source":

                            file.name,

                            "cell":

                            idx

                        }

                    )

                )


        return docs


    def load_zip(

        self,

        file

    ):

        docs = []


        with zipfile.ZipFile(

            file,

            "r"

        ) as zip_ref:


            for filename in zip_ref.namelist():


                if filename.endswith(

                    (

                        ".py",

                        ".ipynb",

                        ".sql",

                        ".txt",

                        ".md",

                        ".json"

                    )

                ):


                    content = zip_ref.read(

                        filename

                    )


                    docs.append(

                        Document(

                            page_content=content.decode(

                                "utf-8",

                                errors="ignore"

                            ),

                            metadata={

                                "source":

                                filename

                            }

                        )

                    )


        return docs

# ==========================================================


# ============================================================
# 6b. DOCUMENT CLEANER — Limpieza explicita del pipeline RAG
# ============================================================

class DocumentCleaner:
    """
    Limpia documentos antes del chunking.
    - Elimina caracteres nulos y espacios multiples
    - Descarta documentos vacios o muy cortos (<50 chars)
    - Anonimiza datos sensibles basicos (email, IDs)
    """
    def __init__(self, min_chars=50, anonimizar=True):
        self.min_chars  = min_chars
        self.anonimizar = anonimizar

    def clean(self, documents):
        limpios, descartados = [], 0
        for doc in documents:
            texto = self._limpiar(doc.page_content)
            if self.anonimizar:
                texto = self._anonimizar(texto)
            if not texto or len(texto.strip()) < self.min_chars:
                descartados += 1
                continue
            if re.fullmatch(r"[\d\s\-\.]+", texto.strip()):
                descartados += 1
                continue
            doc.page_content = texto
            limpios.append(doc)
        log_debug("DocumentCleaner",
                  f"{len(documents)} entrada -> {len(limpios)} validos, {descartados} descartados")
        return limpios

    def _limpiar(self, texto):
        if not texto:
            return ""
        texto = texto.replace("\x00", " ")
        texto = re.sub(r"[\r\n\t]+", " ", texto)
        texto = re.sub(r" {2,}", " ", texto)
        return texto.strip()

    def _anonimizar(self, texto):
        texto = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[EMAIL]", texto)
        texto = re.sub(r"\b\d{8,12}\b", "[ID_NUMERICO]", texto)
        texto = re.sub(r"\b9\d{8}\b", "[TELEFONO]", texto)
        return texto

# 7. PIPELINE RAG - ADAPTIVE CHUNKER
# ==========================================================

class AdaptiveChunker:


    def split_documents(

        self,

        documents

    ):


        chunks = []


        for doc in documents:


            source = doc.metadata.get(

                "source",

                ""

            ).lower()


            # Documentos metodológicos extensos

            if source.endswith(

                (

                    ".pdf",

                    ".docx"

                )

            ):


                splitter = RecursiveCharacterTextSplitter(

                    chunk_size=1800,

                    chunk_overlap=350

                )


            # Documentos de texto


            elif source.endswith(

                (

                    ".txt",

                    ".md"

                )

            ):


                splitter = RecursiveCharacterTextSplitter(

                    chunk_size=1500,

                    chunk_overlap=250

                )


            # Código y notebooks


            elif source.endswith(

                (

                    ".py",

                    ".sql",

                    ".ipynb",

                    ".json"

                )

            ):


                splitter = RecursiveCharacterTextSplitter(

                    chunk_size=1200,

                    chunk_overlap=150

                )


            # Métricas y datasets


            elif source.endswith(

                (

                    ".csv",

                    ".xlsx",

                    ".xls"

                )

            ):


                splitter = RecursiveCharacterTextSplitter(

                    chunk_size=1500,

                    chunk_overlap=300

                )


            # Otros formatos


            else:


                splitter = RecursiveCharacterTextSplitter(

                    chunk_size=1000,

                    chunk_overlap=200

                )


            chunks.extend(

                splitter.split_documents(

                    [doc]

                )

            )


        return chunks

# ==========================================================


# ==========================================================
# 8. PIPELINE RAG - METADATA MANAGER (con normativa SBS)
# ==========================================================

class MetadataManager:
    NORMATIVA_MAP = {
        "3780": "SBS-3780-2011",
        "00053": "SBS-00053-2023",
        "272":  "SBS-272-2017",
        "1445": "SBS-1445-2021",
    }

    def enrich(self, chunks):
        output = []
        for chunk in chunks:
            meta = chunk.metadata.copy()
            if "normativa_sbs" not in meta:
                meta["normativa_sbs"] = self._detect_normativa(
                    chunk.page_content, meta.get("source",""))
            if "articulo_sbs" not in meta:
                meta["articulo_sbs"] = self._detect_articulo(chunk.page_content)
            if "tipo_documento" not in meta:
                meta["tipo_documento"] = self.detect_type(meta.get("source",""))
            meta.setdefault("caso_id", "RIESGO_CREDITICIO_2026")
            meta.setdefault("version", "v1")
            meta.setdefault("fecha", datetime.today().strftime("%Y-%m-%d"))
            meta.setdefault("topics", [])
            chunk.metadata = meta
            output.append(chunk)
        return output

    def _detect_normativa(self, texto, source=""):
        combined = (texto + " " + source).lower()
        for num, res in self.NORMATIVA_MAP.items():
            if num in combined:
                return res
        if "sbs" in combined:
            return "SBS-general"
        return None

    def _detect_articulo(self, texto):
        m = re.search(r"art[ií]culo\s+(\d+)", texto, re.IGNORECASE)
        return f"Art.{m.group(1)}" if m else None

    def detect_type(self, filename):
        f = filename.lower()
        if any(k in f for k in ["normativa","sbs","resolucion","res_"]):
            return "normativa_sbs"
        if any(k in f for k in ["metodologia","modelo","score"]):
            return "metodologia"
        if any(k in f for k in ["benchmark","metricas","performance"]):
            return "benchmark"
        if any(k in f for k in ["hallazgo"]):
            return "hallazgos"
        if f.endswith((".py",".ipynb",".sql")):
            return "codigo"
        return "general"


# ==========================================================
# 8b. HYBRID INDEXER — Indexa con embeddings en ES
# ==========================================================

class HybridIndexer:
    """Indexa documentos con dense_vector + page_content en Elasticsearch."""
    def __init__(self, es_client, emb_model):
        self.es  = es_client
        self.emb = emb_model

    def indexar(self, index_name, documentos, batch_size=30):
        if self.es is None:
            return 0
        total = 0
        for i in range(0, len(documentos), batch_size):
            batch  = documentos[i:i+batch_size]
            textos = [d.page_content for d in batch]
            try:
                vectores = self.emb.embed_documents(textos) if self.emb else [[0.0]*1536]*len(batch)
            except Exception:
                vectores = [[0.0]*1536] * len(batch)
            for doc, vec in zip(batch, vectores):
                try:
                    self.es.index(
                        index    = index_name,
                        id       = str(uuid.uuid4()),
                        document = {
                            "page_content": doc.page_content,
                            "embedding":    vec,
                            "metadata":     {k: v for k, v in doc.metadata.items()}
                        }
                    )
                    total += 1
                except Exception:
                    pass
        log_debug("HybridIndexer", f"{total} docs indexados en '{index_name}'")
        return total


# ==========================================================
# 8c. HYBRID RETRIEVER — BM25 + kNN + RRF
# ==========================================================

class HybridRetriever:
    """
    Busqueda hibrida en Elasticsearch:
    - BM25: recuperacion lexica exacta (siglas, terminos tecnicos SBS)
    - kNN: recuperacion semantica (conceptos relacionados)
    - RRF: fusion de rankings (ES >= 8.9), fallback manual para < 8.9
    """
    def __init__(self, es_client, index_name, emb_model, k=5):
        self.es         = es_client
        self.index_name = index_name
        self.emb        = emb_model
        self.k          = k
        self._rrf_ok    = check_rrf(es_client) if es_client else False

    def search(self, pregunta, filtro_normativa=None, filtro_tipo=None, k=None):
        k = k or self.k
        if self.es is None or not self.es.indices.exists(index=self.index_name):
            return []
        try:
            qvec = self.emb.embed_query(pregunta) if self.emb else None
        except Exception:
            qvec = None

        filtros = []
        if filtro_normativa:
            filtros.append({"term": {"metadata.normativa_sbs": filtro_normativa}})
        if filtro_tipo:
            filtros.append({"term": {"metadata.tipo_documento": filtro_tipo}})

        if self._rrf_ok and qvec:
            return self._rrf(pregunta, qvec, k, filtros)
        elif qvec:
            return self._hybrid_manual(pregunta, qvec, k, filtros)
        return self._bm25(pregunta, k, filtros)

    def _rrf(self, p, qvec, k, filtros):
        body = {"retriever":{"rrf":{
            "retrievers":[
                {"standard":{"query":{"match":{"page_content":p}}}},
                {"knn":{"field":"embedding","query_vector":qvec,"num_candidates":k*4}}
            ],
            "rank_window_size":k*3,"rank_constant":60
        }},"size":k}
        if filtros:
            body["retriever"]["rrf"]["filter"] = {"bool":{"must":filtros}}
        try:
            return self._parse(self.es.search(index=self.index_name, body=body))
        except Exception:
            return self._hybrid_manual(p, qvec, k, filtros)

    def _hybrid_manual(self, p, qvec, k, filtros):
        fc = {"bool":{"must":filtros}} if filtros else {"match_all":{}}
        body = {
            "query":{"bool":{"must":{"match":{"page_content":p}},"filter":fc}},
            "knn":{"field":"embedding","query_vector":qvec,"num_candidates":k*4,"k":k,"filter":fc},
            "size":k
        }
        try:
            return self._parse(self.es.search(index=self.index_name, body=body))
        except Exception:
            return self._bm25(p, k, filtros)

    def _bm25(self, p, k, filtros):
        q = ({"bool":{"must":{"match":{"page_content":p}},"filter":filtros}}
             if filtros else {"match":{"page_content":p}})
        try:
            return self._parse(self.es.search(index=self.index_name, body={"query":q,"size":k}))
        except Exception:
            return []

    def _parse(self, resp):
        return [{"page_content":h["_source"].get("page_content",""),
                 "metadata":h["_source"].get("metadata",{}),"score":round(h.get("_score") or 0,4)}
                for h in resp["hits"]["hits"]]


# ==========================================================
# 8d. HyDE — Hypothetical Document Embeddings
# ==========================================================

HYDE_SYSTEM_PROMPT = (
    "Eres un experto en validacion de modelos de riesgo crediticio "
    "y normativa bancaria peruana (SBS). Genera un fragmento tecnico de 2-4 oraciones "
    "que responderia directamente la pregunta, como si fuera extraido de un documento "
    "oficial SBS o informe de validacion. Solo el fragmento, sin introduccion."
)

class HyDERetriever:
    """
    Hypothetical Document Embeddings.

    Por que mejora el recall:
      La pregunta corta tiene un vector distante de los chunks tecnicos indexados.
      HyDE pide al LLM que genere un documento hipotetico que simularia la respuesta
      real, luego vectoriza ese documento (no la pregunta). El vector resultante es
      mucho mas cercano a los chunks relevantes indexados.

    Flujo:
      pregunta -> LLM genera doc hipotetico -> embed(doc) -> kNN + BM25 -> chunks reales
    """
    def __init__(self, base_retriever, emb_model, use_hyde=True):
        self.base   = base_retriever
        self.emb    = emb_model
        self._cache = {}
        api_key     = os.environ.get("OPENAI_API_KEY","")
        self.llm    = (ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
                       if (api_key and use_hyde) else None)

    def _doc_hipotetico(self, pregunta):
        if pregunta in self._cache:
            return self._cache[pregunta]
        if not self.llm:
            return pregunta
        try:
            doc = self.llm.invoke([
                SystemMessage(content=HYDE_SYSTEM_PROMPT),
                HumanMessage(content=f"Pregunta: {pregunta}\n\nFragmento hipotetico:")
            ]).content.strip()
            self._cache[pregunta] = doc
            return doc
        except Exception:
            return pregunta

    def search(self, pregunta, filtro_normativa=None, filtro_tipo=None, k=5):
        if not self.llm:
            return self.base.search(pregunta, filtro_normativa, filtro_tipo, k)

        doc_h = self._doc_hipotetico(pregunta)
        try:
            hvec = self.emb.embed_query(doc_h) if self.emb else None
        except Exception:
            hvec = None

        if hvec is None:
            return self.base.search(pregunta, filtro_normativa, filtro_tipo, k)

        base = self.base
        if base.es is None or not base.es.indices.exists(index=base.index_name):
            return []

        filtros = []
        if filtro_normativa:
            filtros.append({"term": {"metadata.normativa_sbs": filtro_normativa}})
        if filtro_tipo:
            filtros.append({"term": {"metadata.tipo_documento": filtro_tipo}})

        if base._rrf_ok:
            body = {"retriever":{"rrf":{
                "retrievers":[
                    {"standard":{"query":{"match":{"page_content":pregunta}}}},
                    {"knn":{"field":"embedding","query_vector":hvec,"num_candidates":k*4}}
                ],
                "rank_window_size":k*3,"rank_constant":60
            }},"size":k}
            if filtros:
                body["retriever"]["rrf"]["filter"] = {"bool":{"must":filtros}}
        else:
            fc = {"bool":{"must":filtros}} if filtros else {"match_all":{}}
            body = {
                "query":{"bool":{"must":{"match":{"page_content":pregunta}},"filter":fc}},
                "knn":{"field":"embedding","query_vector":hvec,"num_candidates":k*4,"k":k,"filter":fc},
                "size":k
            }
        try:
            resp = base.es.search(index=base.index_name, body=body)
            return [{"page_content":h["_source"].get("page_content",""),
                     "metadata":h["_source"].get("metadata",{}),"score":round(h.get("_score") or 0,4),
                     "hyde_doc":doc_h}
                    for h in resp["hits"]["hits"]]
        except Exception:
            return self.base.search(pregunta, filtro_normativa, filtro_tipo, k)

    def search_trazado(self, pregunta, filtro_normativa=None, k=5):
        doc_h = self._doc_hipotetico(pregunta) if self.llm else "HyDE desactivado"
        return {"pregunta":pregunta,"doc_hipotetico":doc_h,
                "resultados":self.search(pregunta, filtro_normativa, k=k)}



# ==========================================================
# 9b. MEMORY MANAGER - puente temporal <-> permanente
# ==========================================================

class MemoryManager:
    """
    Capa unificada de memoria hibrida.

    Temporal   -> st.session_state  (rapida, vive en sesion activa)
    Permanente -> SQLite            (persiste entre sesiones)

    Reglas:
      - inicializar_session(): SQLite -> session_state al arrancar
      - Durante sesion: agente lee de session_state sin tocar disco
      - persistir_hallazgo(): session_state -> SQLite solo tras HITL
      - NUNCA escribe a SQLite sin decision humana explicita
    """

    def __init__(self, db_name=None, limite_contexto=5):
        self.db     = db_name or DB_NAME
        self.limite = limite_contexto

    def cargar_contexto_historico(self, keyword: str) -> str:
        """Lee hallazgos previos de SQLite. Solo lectura."""
        try:
            conn = sqlite3.connect(self.db)
            df   = pd.read_sql_query(
                """SELECT caso_id, severidad, hallazgo, fecha
                   FROM hallazgos WHERE hallazgo LIKE ?
                   ORDER BY fecha DESC LIMIT ?""",
                conn, params=(f"%{keyword}%", self.limite)
            )
            conn.close()
            return df.to_string(index=False) if not df.empty else ""
        except Exception:
            return ""

    def inicializar_session(self, modelo: str, periodo: str) -> str:
        """
        Al arrancar una sesion de validacion:
        1. Carga hallazgos previos del modelo desde SQLite
        2. Los inyecta en session_state como contexto historico
        3. Inicializa claves de UI si no existen aun
        Retorna el historico como string para el agente.
        """
        for clave, valor in {
            "reporte_preliminar":      "",
            "evidencia_rag":           "",
            "ultimo_prompt":           "",
            "ultimo_resumen_metricas": "",
            "ultimo_error":            "",
        }.items():
            if clave not in st.session_state:
                st.session_state[clave] = valor

        clave_hist = f"memoria_historica_{modelo}"
        if clave_hist not in st.session_state:
            historico = self.cargar_contexto_historico(f"{modelo}_{periodo}")
            st.session_state[clave_hist] = historico

        return st.session_state.get(clave_hist, "")

    def persistir_hallazgo(
        self, modelo, periodo, categoria, hallazgo,
        severidad, impacto, recomendacion, decision_humana, comentario
    ) -> bool:
        """
        Escribe a SQLite SOLO si decision_humana != RECHAZADO.
        Actualiza session_state para reflejar el cambio sin reload.
        Retorna True si se persistio, False si fue rechazado.
        """
        if decision_humana == "RECHAZADO":
            log_debug("MemoryManager", "Hallazgo rechazado - no persistido", "OK")
            return False

        caso_id    = f"{modelo}_{periodo}"
        texto_full = (
            f"Categoria: {categoria}\n"
            f"Hallazgo: {hallazgo}\n"
            f"Impacto: {impacto}\n"
            f"Recomendacion: {recomendacion}\n"
            f"Decision: {decision_humana}\n"
            f"Comentario validador: {comentario}"
        )
        try:
            conn = sqlite3.connect(self.db)
            conn.execute(
                "INSERT INTO hallazgos (caso_id, hallazgo, severidad) VALUES (?,?,?)",
                (caso_id, texto_full, severidad)
            )
            conn.commit()
            conn.close()
            # Reflejar en session_state sin reload
            clave_hist = f"memoria_historica_{modelo}"
            prev  = st.session_state.get(clave_hist, "")
            entry = f"\n[{datetime.now().strftime('%Y-%m-%d')}] {severidad}: {hallazgo[:100]}..."
            st.session_state[clave_hist] = prev + entry
            log_debug("MemoryManager", f"Persistido [{decision_humana}] sev:{severidad}", "OK")
            actualizar_estado_flujo("memoria", "OK")
            return True
        except Exception as e:
            log_debug("MemoryManager", f"Error: {e}", "Error")
            return False

    def conteo(self, tabla: str) -> int:
        """Cuenta filas en una tabla SQLite."""
        try:
            conn = sqlite3.connect(self.db)
            n = conn.execute(f"SELECT COUNT(*) FROM {tabla}").fetchone()[0]
            conn.close()
            return n
        except Exception:
            return 0

    def leer_tabla(self, tabla: str) -> pd.DataFrame:
        """Lee tabla SQLite para visualizacion en UI."""
        if tabla not in ["hallazgos", "ejecuciones", "rag_traces"]:
            return pd.DataFrame()
        try:
            conn = sqlite3.connect(self.db)
            df   = pd.read_sql_query(
                f"SELECT * FROM {tabla} ORDER BY fecha DESC LIMIT 50", conn)
            conn.close()
            return df
        except Exception:
            return pd.DataFrame()


memory_manager = MemoryManager(limite_contexto=5)

# 9. MEMORIA SQLITE
# ==========================================================

# DB_NAME definido en seccion 1 - no redefinir


def inicializar_memoria():

    conn = sqlite3.connect(DB_NAME)

    cursor = conn.cursor()


    cursor.execute("""

    CREATE TABLE IF NOT EXISTS hallazgos (

        id INTEGER PRIMARY KEY AUTOINCREMENT,

        caso_id TEXT,

        hallazgo TEXT,

        severidad TEXT,

        fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP

    )

    """)


    cursor.execute("""

    CREATE TABLE IF NOT EXISTS ejecuciones (

        id INTEGER PRIMARY KEY AUTOINCREMENT,

        modelo TEXT,

        periodo TEXT,

        estado TEXT,

        resumen TEXT,

        error TEXT,

        fecha TIMESTAMP DEFAULT CURRENT_TIMESTAMP

    )

    """)


    conn.commit()

    conn.close()


def guardar_hallazgo(

    caso_id,

    hallazgo,

    severidad

):

    conn = sqlite3.connect(DB_NAME)

    cursor = conn.cursor()


    cursor.execute(

        """

        INSERT INTO hallazgos (

            caso_id,

            hallazgo,

            severidad

        )

        VALUES (

            ?, ?, ?

        )

        """,

        (

            caso_id,

            hallazgo,

            severidad

        )

    )


    conn.commit()

    conn.close()


def guardar_ejecucion(

    modelo,

    periodo,

    estado,

    resumen,

    error=""

):

    conn = sqlite3.connect(DB_NAME)

    cursor = conn.cursor()


    cursor.execute(

        """

        INSERT INTO ejecuciones (

            modelo,

            periodo,

            estado,

            resumen,

            error

        )

        VALUES (

            ?, ?, ?, ?, ?

        )

        """,

        (

            modelo,

            periodo,

            estado,

            resumen,

            error

        )

    )


    conn.commit()

    conn.close()


def consultar_hallazgos_previos(

    keyword,

    limite=5

):

    conn = sqlite3.connect(DB_NAME)


    query = """

    SELECT *

    FROM hallazgos

    WHERE hallazgo LIKE ?

    ORDER BY fecha DESC

    LIMIT ?

    """


    df = pd.read_sql_query(

        query,

        conn,

        params=(

            f"%{keyword}%",

            limite

        )

    )


    conn.close()


    if df.empty:

        return "No se encontraron hallazgos históricos."


    return df.to_string(

        index=False

    )


def guardar_hallazgo_validado(
    modelo, periodo, categoria, hallazgo,
    severidad, impacto, recomendacion,
    decision_humana, comentario_humano
):
    """Usa MemoryManager: persiste en SQLite Y actualiza session_state."""
    return memory_manager.persistir_hallazgo(
        modelo          = modelo,
        periodo         = periodo,
        categoria       = categoria,
        hallazgo        = hallazgo,
        severidad       = severidad,
        impacto         = impacto,
        recomendacion   = recomendacion,
        decision_humana = decision_humana,
        comentario      = comentario_humano
    )

def obtener_tabla(

    tabla

):


    tablas_permitidas = [

        "hallazgos",

        "ejecuciones"

    ]


    if tabla not in tablas_permitidas:

        return pd.DataFrame()


    try:

        conn = sqlite3.connect(

            DB_NAME

        )


        query = f"""

            SELECT *

            FROM {tabla}

            ORDER BY fecha DESC

            """



        df = pd.read_sql_query(

            query,

            conn

        )

        conn.close()


        return df


    except Exception:

        return pd.DataFrame()


# ============================================================

# 10. MÉTRICAS
# ============================================================

def calcular_gini(
    y_true,
    y_score
) -> Optional[float]:

    try:

        auc = roc_auc_score(
            y_true,
            y_score
        )

        return float(
            2 * auc - 1
        )

    except Exception:

        return None


def dataframe_to_json_records(
    df: pd.DataFrame
) -> str:

    return df.to_json(
        orient="records",
        force_ascii=False
    )


def json_records_to_dataframe(
    df_json: str
) -> pd.DataFrame:

    data = safe_json_loads(
        df_json,
        default=[]
    )

    if not data:

        return pd.DataFrame()

    return pd.DataFrame(
        data
    )


def resumir_excel_gini(
    df: pd.DataFrame
) -> str:

    try:

        col_gini = None

        for col in df.columns:

            if "gini" in col.lower():

                col_gini = col

                break


        if col_gini is None:

            return ""


        col_periodo = None

        for col in df.columns:

            if col.lower() == "periodo":

                col_periodo = col

                break


        df_aux = df.copy()


        df_aux[col_gini] = pd.to_numeric(

            df_aux[col_gini],

            errors="coerce"

        )


        df_aux = df_aux.dropna(

            subset=[col_gini]

        )


        if df_aux.empty:

            return ""


        gini_inicio = float(

            df_aux[col_gini].iloc[0]

        )


        gini_fin = float(

            df_aux[col_gini].iloc[-1]

        )


        variacion = (

            gini_fin - gini_inicio

        )


        if col_periodo:

            periodo_inicio = str(

                df_aux[col_periodo].iloc[0]

            )


            periodo_fin = str(

                df_aux[col_periodo].iloc[-1]

            )

        else:

            periodo_inicio = "inicio"

            periodo_fin = "fin"


        if variacion < 0:

            tendencia = "decreciente"

        elif variacion > 0:

            tendencia = "creciente"

        else:

            tendencia = "estable"


        resumen = {

            "tipo_archivo":

            "serie_metricas",


            "periodo_inicio":

            periodo_inicio,


            "periodo_fin":

            periodo_fin,


            "gini_inicio":

            round(

                gini_inicio,

                4

            ),


            "gini_fin":

            round(

                gini_fin,

                4

            ),


            "variacion_gini":

            round(

                variacion,

                4

            ),


            "tendencia_gini":

            tendencia,


            "observacion":

            (

                "El Gini muestra tendencia decreciente; podría requerir análisis de degradación de performance."

                if variacion < 0

                else

                "El Gini no muestra caída en la serie revisada."

            )

        }


        return json.dumps(

            resumen,

            ensure_ascii=False,

            indent=2

        )


    except Exception as e:

        return (

            f"No se pudo resumir tendencia de Gini: {str(e)}"

        )


# ============================================================

# 11. BASEMODEL INPUTS
# ============================================================


class AnalizarDocumentoInput(BaseModel):
    texto_metodologia: str = Field(
        description="Texto extraído del documento metodológico del modelo de riesgo."
    )


class AnalizarCodigoInput(BaseModel):
    codigo: str = Field(
        description="Código Python, SQL o notebook asociado al modelo de riesgo."
    )


class CalcularMetricasInput(BaseModel):
    df_json: str = Field(
        description="Dataset en formato JSON records."
    )


class ConsultarMemoriaInput(BaseModel):
    keyword: str = Field(
        description="Palabra clave para buscar hallazgos históricos en memoria SQLite."
    )


class BenchmarkInput(BaseModel):
    descripcion_hallazgo: str = Field(
        description="Descripción del hallazgo preliminar."
    )


class EvaluarHallazgoInput(BaseModel):
    hallazgo: str = Field(
        description="Hallazgo preliminar identificado por el agente o por el validador."
    )


# ============================================================

# 12. TOOLS
# ============================================================

@tool(args_schema=AnalizarDocumentoInput)
def analizar_documento_metodologico(texto_metodologia: str) -> str:
    """
    Analiza el documento metodológico del modelo de riesgo e identifica elementos mínimos, faltantes y posibles alertas.
    """
    try:
        if not texto_metodologia or len(texto_metodologia.strip()) < 100:
            return fallback_falta_info("documento metodológico")

        texto = texto_metodologia.lower()

        checks = {
            "objetivo_modelo": any(x in texto for x in ["objetivo", "propósito", "proposito"]),
            "poblacion_objetivo": any(x in texto for x in ["población objetivo", "poblacion objetivo", "universo", "muestra"]),
            "variables": any(x in texto for x in ["variables", "feature", "predictor"]),
            "metodologia": any(x in texto for x in ["metodología", "metodologia", "modelo", "algoritmo", "regresión", "regresion"]),
            "metricas": any(x in texto for x in ["gini", "auc", "ks", "psi", "performance", "calibración", "calibracion"]),
            "limitaciones": any(x in texto for x in ["limitación", "limitacion", "supuesto", "restricción", "restriccion"]),
            "criterios_exclusion": any(x in texto for x in ["excluyen", "exclusión", "exclusion", "filtro", "score nulo"])
        }

        faltantes = [k for k, v in checks.items() if not v]

        salida = {
            "checks_documento": checks,
            "posibles_faltantes": faltantes,
            "observacion": (
                "El documento contiene elementos mínimos esperados."
                if not faltantes
                else "El documento podría requerir mayor detalle en: " + ", ".join(faltantes)
            )
        }

        return json.dumps(salida, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"[ERROR_TOOL_DOCUMENTO] {str(e)}"


@tool(args_schema=AnalizarCodigoInput)
def analizar_codigo_modelo(codigo: str) -> str:
    """
    Analiza el código del modelo de riesgo para revisar población, filtros, métricas, outputs y trazabilidad.
    """
    try:
        if not codigo or len(codigo.strip()) < 50:
            return fallback_falta_info("código del modelo")

        c = codigo.lower()

        checks = {
            "usa_sql_o_athena": any(x in c for x in ["select ", "from ", "join ", "athena", "awswrangler", "wr.athena", "pd.read_sql"]),
            "construye_poblacion": any(x in c for x in ["poblacion", "población", "universo", "base", "target", "preparar_poblacion"]),
            "aplica_filtros": any(x in c for x in ["where", "filter", "query", "dropna", "isin", "fecha", "periodo", "notna", "mora_previa"]),
            "calcula_metricas": any(x in c for x in ["gini", "auc", "roc_auc", "ks", "psi", "performance", "roc_auc_score"]),
            "genera_outputs": any(x in c for x in ["to_csv", "to_parquet", "to_excel", "s3", "parquet", "xlsx"]),
            "manejo_errores": any(x in c for x in ["try:", "except", "raise"]),
            "benchmark": any(x in c for x in ["benchmark", "train_test_split", "logisticregression", "pipeline"]),
            "target": any(x in c for x in ["target", "mora_futura", "incumplimiento"])
        }

        faltantes = [k for k, v in checks.items() if not v]

        salida = {
            "checks_codigo": checks,
            "posibles_faltantes": faltantes,
            "observacion": (
                "El código evidencia pasos relevantes para validación."
                if not faltantes
                else "El código podría requerir revisión en: " + ", ".join(faltantes)
            )
        }

        return json.dumps(salida, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"[ERROR_TOOL_CODIGO] {str(e)}"


@tool(args_schema=CalcularMetricasInput)
def calcular_metricas_validacion(df_json: str) -> str:
    """
    Calcula métricas básicas de validación como tasa de target, Gini, tendencia de Gini y resumen de columnas.
    """
    try:
        df = json_records_to_dataframe(df_json)

        if df.empty:
            return fallback_falta_info("dataset para métricas")

        resumen = {
            "n_registros": int(len(df)),
            "columnas": list(df.columns),
            "metricas": {}
        }

        if "target" in df.columns:
            resumen["metricas"]["tasa_target"] = float(
                pd.to_numeric(df["target"], errors="coerce").mean()
            )

        if "score" in df.columns and "target" in df.columns:
            y = pd.to_numeric(df["target"], errors="coerce")
            s = pd.to_numeric(df["score"], errors="coerce")

            valid = y.notna() & s.notna()

            if valid.sum() > 2 and y[valid].nunique() == 2:
                resumen["metricas"]["gini"] = calcular_gini(y[valid], s[valid])
            else:
                resumen["metricas"]["gini"] = "No calculable: target debe tener dos clases y suficientes registros."

        #cols_lower = {c.lower(): c for c in df.columns}

        #if "gini" in cols_lower:
        #    col_gini = cols_lower["gini"]
        col_gini = None
        for col in df.columns:
            if "gini" in col.lower():
                col_gini = col
                break
        if col_gini:
            g = pd.to_numeric(df[col_gini], errors="coerce").dropna()

            if len(g) >= 2:
                resumen["metricas"]["gini_inicio"] = float(g.iloc[0])
                resumen["metricas"]["gini_fin"] = float(g.iloc[-1])
                resumen["metricas"]["variacion_gini"] = float(g.iloc[-1] - g.iloc[0])
                resumen["metricas"]["tendencia_gini"] = (
                    "decreciente"
                    if g.iloc[-1] < g.iloc[0]
                    else "creciente"
                    if g.iloc[-1] > g.iloc[0]
                    else "estable"
                )

        resumen["metricas"]["columnas_numericas"] = df.select_dtypes(include=[np.number]).columns.tolist()

        return json.dumps(resumen, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"[ERROR_TOOL_METRICAS] {str(e)}"


@tool(args_schema=ConsultarMemoriaInput)
def consultar_memoria_tool(keyword: str) -> str:
    """
    Consulta hallazgos historicos validados en la memoria SQLite,
    filtrando por una palabra clave relacionada con el caso o modelo.
    """
    try:
        return consultar_hallazgos_previos(keyword=keyword, limite=5)
    except Exception as e:
        return f"[ERROR_TOOL_MEMORIA] {str(e)}"

@tool(args_schema=BenchmarkInput)
def ejecutar_benchmark_tool(descripcion_hallazgo: str) -> str:
    """
    Sugiere escenarios de benchmark o análisis alternativo cuando un hallazgo puede afectar resultados del modelo.
    """
    try:
        if not descripcion_hallazgo or len(descripcion_hallazgo.strip()) < 20:
            return fallback_falta_info("descripción del hallazgo para benchmark")

        texto = descripcion_hallazgo.lower()
        escenarios = []

        if any(x in texto for x in ["población", "poblacion", "filtro", "exclusión", "exclusion"]):
            escenarios.append("Comparar resultados con y sin el filtro observado en la población objetivo.")

        if any(x in texto for x in ["variable", "correlación", "correlacion", "shap", "gain"]):
            escenarios.append("Evaluar estabilidad del modelo excluyendo o reemplazando variables observadas.")

        if any(x in texto for x in ["gini", "auc", "ks", "performance", "degradación", "degradacion"]):
            escenarios.append("Comparar métricas de performance entre metodología original y alternativa.")

        if any(x in texto for x in ["psi", "estabilidad", "drift"]):
            escenarios.append("Comparar estabilidad de distribución por periodos y segmentos.")

        if any(x in texto for x in ["target", "mora", "incumplimiento"]):
            escenarios.append("Validar sensibilidad ante cambios en la definición de target o ventana de observación.")

        if not escenarios:
            escenarios.append("Definir escenario alternativo documentando supuesto, cambio aplicado e impacto esperado.")

        salida = {
            "benchmark_recomendado": escenarios,
            "nota": "El benchmark debe ser revisado por el validador antes de considerarse conclusión."
        }

        return json.dumps(salida, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"[ERROR_TOOL_BENCHMARK] {str(e)}"


@tool(args_schema=EvaluarHallazgoInput)
def evaluar_hallazgo_tool(hallazgo: str) -> str:
    """
    Evalúa un hallazgo preliminar y sugiere severidad, impacto y necesidad de benchmark.
    """
    try:
        if not hallazgo or len(hallazgo.strip()) < 20:
            return fallback_falta_info("hallazgo para evaluación")

        h = hallazgo.lower()

        impacto_alto_keywords = [
            "población objetivo", "poblacion objetivo", "target", "filtro",
            "exclusión", "exclusion", "gini", "auc", "ks", "calibración",
            "calibracion", "degradación", "degradacion", "performance"
        ]

        impacto_medio_keywords = [
            "documentación", "documentacion", "variable", "supuesto",
            "trazabilidad", "reproducibilidad", "código", "codigo",
            "psi", "estabilidad", "drift"
        ]

        if any(k in h for k in impacto_alto_keywords):
            severidad = "Alta"
            impacto = "Puede afectar resultados, población o interpretación del modelo."
        elif any(k in h for k in impacto_medio_keywords):
            severidad = "Media"
            impacto = "Puede afectar claridad metodológica, trazabilidad, estabilidad o reproducibilidad."
        else:
            severidad = "Baja"
            impacto = "Impacto acotado o principalmente documental."

        salida = {
            "severidad_sugerida": severidad,
            "impacto_sugerido": impacto,
            "requiere_benchmark": severidad in ["Alta", "Media"],
            "nota": "Clasificación sugerida; debe ser validada por el humano."
        }

        return json.dumps(salida, ensure_ascii=False, indent=2)

    except Exception as e:
        return f"[ERROR_TOOL_EVALUACION] {str(e)}"
# ==========================================================



# ==========================================================
# TOOL RAG — buscar_evidencia_rag con HyDE integrado
# ==========================================================

class BuscarEvidenciaInput(BaseModel):
    pregunta:         str  = Field(description="Pregunta o tema a buscar en la base de conocimiento.")
    filtro_normativa: str  = Field(default=None, description="Resolucion SBS (ej: SBS-3780-2011).")
    filtro_tipo:      str  = Field(default=None, description="Tipo: normativa_sbs, metodologia, benchmark, codigo.")
    usar_hyde:        bool = Field(default=True, description="Activar HyDE para mejor recall semantico.")

@tool(args_schema=BuscarEvidenciaInput)
def buscar_evidencia_rag(
    pregunta:         str,
    filtro_normativa: str  = None,
    filtro_tipo:      str  = None,
    usar_hyde:        bool = True
) -> str:
    """
    Busca evidencia documental usando RAG hibrido con HyDE.
    HyDE genera primero un documento hipotetico para mejorar el recall semantico.
    Retorna fragmentos con: fuente, pagina, normativa SBS, articulo y score.
    """
    es_cl   = init_elasticsearch()
    emb_cl  = init_embeddings_model()
    base_r  = HybridRetriever(es_cl, INDEX_NAME, emb_cl, k=5) if es_cl else None
    hyde_r  = HyDERetriever(base_r, emb_cl, use_hyde=True) if base_r else None

    r_activo, modo = None, "sin_retriever"
    if usar_hyde and hyde_r is not None and hyde_r.llm is not None:
        r_activo, modo = hyde_r, "HyDE"
    elif base_r is not None:
        r_activo, modo = base_r, "Hibrido"

    if r_activo is None:
        return "Retriever no disponible. Configura ELASTIC_URL, ELASTIC_API_KEY y OPENAI_API_KEY en api.txt."

    t0 = time.time()
    hyde_doc_log = None
    try:
        resultados = r_activo.search(pregunta, filtro_normativa, filtro_tipo, k=5)
        if modo == "HyDE" and hasattr(r_activo, "_doc_hipotetico"):
            hyde_doc_log = r_activo._doc_hipotetico(pregunta)
    except Exception as e:
        log_rag_trace(pregunta, None, "error", [], int((time.time()-t0)*1000), str(e))
        return f"Error RAG: {e}"

    log_rag_trace(pregunta, hyde_doc_log, modo, resultados, int((time.time()-t0)*1000))
    actualizar_estado_flujo("rag", "OK")

    if not resultados:
        return (f"No se encontro evidencia en la base de conocimiento [modo: {modo}].\n"
                "Considera ampliar la base de conocimiento o reformular la pregunta.")

    lineas = [f"Evidencia — {len(resultados)} fragmentos [modo: {modo}]:\n"]
    for i, doc in enumerate(resultados, 1):
        meta   = doc.get("metadata", {})
        header = (f"[{i}] Fuente: {meta.get('source','?')} | Pag: {meta.get('page','—')} "
                  f"| Score: {doc.get('score','—')}")
        if meta.get("normativa_sbs"):
            header += f" | {meta['normativa_sbs']}"
        if meta.get("articulo_sbs"):
            header += f" | {meta['articulo_sbs']}"
        if doc.get("hyde_doc") and i == 1:
            lineas.append(f"[HyDE doc hipotetico]: {doc['hyde_doc'][:200]}...")
        lineas.append(header)
        lineas.append(doc.get("page_content","")[:900])
        lineas.append("")
    return "\n".join(lineas)

TOOLS = [
    analizar_documento_metodologico,
    analizar_codigo_modelo,
    calcular_metricas_validacion,
    consultar_memoria_tool,
    ejecutar_benchmark_tool,
    evaluar_hallazgo_tool,
    buscar_evidencia_rag,
]

# 13. AGENTE
# ============================================================

SYSTEM_PROMPT = """
Eres ValidAI Risk, un asistente experto en validación de modelos de riesgo crediticio en banca.

Tu objetivo es apoyar al equipo de validación de modelos en:
1. Revisar el documento metodológico.
2. Revisar código Python, SQL o notebooks asociados al modelo.
3. Contrastar metodología versus implementación.
4. Identificar hallazgos metodológicos, técnicos y de trazabilidad.
5. Evaluar severidad, impacto y recurrencia.
6. Recomendar benchmark o análisis alternativos cuando el hallazgo pueda afectar resultados.
7. Generar un reporte preliminar para revisión humana.

Reglas obligatorias:
- No reemplazas al validador humano.
- No inventes información si no está en los insumos.
- Si falta información, dilo explícitamente.
- Diferencia entre hallazgo, riesgo, impacto y recomendación.
- Mantén lenguaje formal, claro y orientado a banca.
- Si la consulta está fuera del dominio de validación de modelos de riesgo, responde que está fuera de alcance.
- No guardes memoria por tu cuenta; el feedback se guarda solo después de revisión humana.
- Si detectas posible impacto en población, metodología, target, filtros, variables o métricas, sugiere benchmark.
- Las alertas se detectan leyendo y analizando la observación, evaluando si tiene impacto en el proyecto o resultado.
- No limites las alertas únicamente a métricas como Gini o PSI.
- Las métricas tabulares son complementarias. Si no existen, no concluyas que el análisis falló.
- La salida debe ser estructurada.

Formato de respuesta:
1. Resumen ejecutivo
2. Evidencia revisada
3. Hallazgos metodológicos
4. Hallazgos de código / implementación
5. Consistencia metodología vs código
6. Evaluación de impacto
7. Benchmark sugerido si aplica
8. Recomendaciones
9. Limitaciones
10. Próximos pasos para revisión humana

"""
REGLAS_RAG = """

REGLAS PARA EL USO DE EVIDENCIA DOCUMENTAL

1. Cuando necesites evidencia documental utiliza buscar_evidencia_rag().

2. Fundamenta la respuesta utilizando únicamente la evidencia recuperada.

3. Siempre cita la fuente documental.

4. Si no existe evidencia suficiente, indícalo explícitamente.

5. Nunca inventes normativas, benchmarks o métricas.

6. Prioriza documentos con normativa_sbs sobre otros tipos.
7. Cita siempre: documento, pagina, resolucion SBS y articulo cuando aplique.
8. Si la respuesta menciona umbrales (Gini, PSI, KS), verifica contra normativa SBS recuperada.

"""

SYSTEM_PROMPT += REGLAS_RAG

def create_validation_agent():
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)

    if AGENT_BACKEND == "langchain_create_agent" and create_agent is not None:
        try:
            return create_agent(
                model=llm,
                tools=TOOLS,
                system_prompt=SYSTEM_PROMPT
            )
        except TypeError:
            return create_agent(
                model=llm,
                tools=TOOLS,
                prompt=SYSTEM_PROMPT
            )

    return create_react_agent(
        model=llm,
        tools=TOOLS,
        prompt=SYSTEM_PROMPT
    )


def ejecutar_agente(agent, prompt_usuario: str) -> str:
    try:
        response = agent.invoke({
            "messages": [
                {
                    "role": "user",
                    "content": prompt_usuario
                }
            ]
        })

        if (
            isinstance(response, dict)
            and "messages" in response
            and len(response["messages"]) > 0
        ):

            return response["messages"][-1].content

        return str(response)

    except Exception as e:
        return (
            "No pude completar el análisis con el agente.\n\n"
            f"Error técnico: {str(e)}\n\n"
            "Fallback: revisa que la API key esté configurada, que los insumos no estén vacíos "
            "y que el modelo tenga acceso a las tools."
        )
# ==========================================================
# CONTEXTO RAG
# ==========================================================

def obtener_contexto_rag(

    pregunta

):
    if not pregunta:

        return ""

    try:


        evidencia = buscar_evidencia_rag.invoke(

            {

               "pregunta": pregunta

            }

        )


        return evidencia


    except Exception:


        return ""

def generar_prompt_validacion(
    modelo,
    periodo,
    texto_metodologia,
    codigo_modelo,
    resumen_metricas,
    memoria_historica,
    observacion_usuario
) -> str:

    consulta_rag = (

    observacion_usuario

    if observacion_usuario

    else modelo

    )


    contexto_rag = obtener_contexto_rag(

        consulta_rag

    )


    return f"""
Analiza el siguiente caso de validación de modelo de riesgo.

MODELO:
{modelo}

PERIODO:
{periodo}

OBSERVACIÓN O FOCO DEL VALIDADOR:
{observacion_usuario if observacion_usuario else "No se indicó una observación específica. Realiza revisión general."}

DOCUMENTO METODOLÓGICO:
{texto_metodologia if texto_metodologia else "No se adjuntó documento metodológico."}

CÓDIGO / NOTEBOOK:
{codigo_modelo if codigo_modelo else "No se adjuntó código."}

RESULTADOS / MÉTRICAS DISPONIBLES:
{resumen_metricas if resumen_metricas else "No se adjuntaron métricas tabulares. Este insumo es complementario y queda pendiente."}

MEMORIA HISTÓRICA CONSULTADA:
{memoria_historica if memoria_historica else "No hay memoria histórica disponible."}

EVIDENCIA DOCUMENTAL (RAG):
{contexto_rag if contexto_rag else "No se encontró evidencia documental adicional."}


TAREA:
Genera un reporte preliminar de validación con enfoque bancario.

Debes:
- Identificar hallazgos metodológicos.
- Identificar hallazgos de código o implementación.
- Contrastar metodología vs código.
- Evaluar si la observación puede tener impacto en el proyecto o en los resultados.
- No limitar el análisis únicamente a métricas como Gini o PSI.
- Si hay impacto, proponer benchmark o análisis alternativo.
- Indicar limitaciones si falta información.
- Dejar claro que la decisión final es del validador humano.
"""


# ============================================================

# 14. INTERFAZ STREAMLIT
# ============================================================

def main():
    configurar_apis_dict()
    inicializar_memoria()
    init_rag_traces_table()
    inicializar_debug()

    # MemoryManager: SQLite -> session_state al arrancar sesion
    modelo_activo  = st.session_state.get("modelo_activo", "Modelo_Riesgo")
    periodo_activo = datetime.now().strftime("%Y-%m")
    _ = memory_manager.inicializar_session(modelo_activo, periodo_activo)

    # Inicializar ES + Embeddings + Retrievers (cacheados por Streamlit)
    es_cl  = init_elasticsearch()
    emb_cl = init_embeddings_model()
    crear_indice_si_no_existe(es_cl)
    base_ret = HybridRetriever(es_cl, INDEX_NAME, emb_cl, k=5) if es_cl else None
    hyde_ret = HyDERetriever(base_ret, emb_cl, use_hyde=True) if base_ret else None

    aplicar_estilo_visual()
    render_logo_header()

    with st.sidebar:
        st.header("Panel de control")
        st.caption("Guía rápida del flujo de validación")

        st.markdown("""
        1. Revisa la arquitectura en **Cómo funciona**.
        2. Carga documentos, código o métricas.
        3. Ejecuta la revisión IA.
        4. Valida el hallazgo humano.
        5. Guarda memoria solo si corresponde.
        """)

        st.divider()

        modelo = st.text_input(
            "Nombre del modelo",
            value="Modelo Riesgo Crediticio"
        )

        periodo = datetime.now().strftime("%Y-%m")

        st.divider()
        st.subheader("Estado del flujo")

        estado_flujo = st.session_state.get("estado_flujo", {})

        def estado_badge(nombre, estado):
            if estado == "OK":
                icono = "✓"
            elif estado == "Error":
                icono = "!"
            elif estado == "En proceso":
                icono = "…"
            else:
                icono = "○"

            st.write(f"{icono} {nombre}: `{estado}`")

        estado_badge("Insumos", estado_flujo.get("insumos", "Pendiente"))
        estado_badge("Preparación", estado_flujo.get("preparacion", "Pendiente"))
        estado_badge("Métricas", estado_flujo.get("metricas", "Pendiente"))
        estado_badge("Memoria", estado_flujo.get("memoria", "Pendiente"))
        estado_badge("RAG", estado_flujo.get("rag", "Pendiente"))
        estado_badge("Evidencia", estado_flujo.get("evidencia", "Pendiente"))
        estado_badge("Agente IA", estado_flujo.get("agente", "Pendiente"))
        estado_badge("Reporte", estado_flujo.get("reporte", "Pendiente"))

        st.divider()
        st.subheader("Estado de memoria")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            st.metric("SQLite hallazgos", memory_manager.conteo("hallazgos"))
        with col_m2:
            clave_h = f"memoria_historica_{st.session_state.get('modelo_activo','')}"
            n_ses = len([l for l in st.session_state.get(clave_h,"").split("\n") if l.strip()])
            st.metric("En sesion", n_ses)
        st.divider()
        st.subheader("Observabilidad")

        tracing = os.environ.get("LANGCHAIN_TRACING_V2", "false")

        if tracing == "true":
            st.success("LangSmith activo")
        else:
            st.info("LangSmith inactivo")

        st.caption("Proyecto: M4-Copiloto-Validacion-Modelos")
        if es_cl:
            st.success("ES conectado")
        else:
            st.warning("ES no configurado")
        if hyde_ret and hyde_ret.llm:
            st.success("HyDE activo")
        else:
            st.info("HyDE inactivo")
        if os.environ.get("LANGCHAIN_TRACING_V2") == "true":
            st.success("LangSmith activo")
        with st.expander("Trazas RAG (HyDE)"):
            df_tr = get_rag_traces(5)
            if not df_tr.empty:
                st.dataframe(df_tr[["timestamp","modo","n_chunks","latencia_ms","pregunta"]].head(5),
                             use_container_width=True)
            else:
                st.caption("Sin trazas aun")

        with st.expander("Ver trazabilidad local"):
            logs = st.session_state.get("debug_logs", [])

            if logs:
                st.dataframe(
                    pd.DataFrame(logs),
                    use_container_width=True,
                    height=220
                )
            else:
                st.caption("Aún no hay eventos registrados.")

            if st.button("Limpiar trazabilidad"):
                limpiar_debug()
                st.success("Trazabilidad limpiada.")

        with st.expander("Último prompt"):
            ultimo_prompt = st.session_state.get("ultimo_prompt", "")

            if ultimo_prompt:
                st.text_area("Prompt enviado", ultimo_prompt, height=220)
            else:
                st.caption("Aún no hay prompt registrado.")

        with st.expander("Últimas métricas"):
            ultimo_resumen = st.session_state.get("ultimo_resumen_metricas", "")

            if ultimo_resumen:
                st.code(ultimo_resumen, language="json")
            else:
                st.caption("Aún no se cargó archivo de métricas.")

        with st.expander("Evidencia recuperada por RAG"):
            evidencia = st.session_state.get("evidencia_rag", "")
            if evidencia:
                st.text(evidencia)
            else:
                st.caption("Aún no hay evidencia recuperada.")


        with st.expander("Errores"):
            ultimo_error = st.session_state.get("ultimo_error", "")

            if ultimo_error:
                st.code(ultimo_error)
            else:
                st.caption("Sin errores registrados.")

    tab1, tab2, tab3, tab4 = st.tabs([
        "🧭 Cómo funciona",
        "🧪 Revisión IA",
        "👤 Validación humana",
        "🗃️ Memoria"
    ])

    with tab1:
        st.subheader("Cómo funciona ValidAI Risk")

        st.markdown("""
        <div style="border: 1px solid #262626; border-radius: 18px; padding: 18px; background: #111111;">
            <span class="green-badge">Arquitectura híbrida</span>
            <h3>Del insumo técnico al hallazgo validado</h3>
            <p>
            ValidAI Risk combina un flujo controlado con un agente IA especializado.
            La solución apoya la revisión metodológica, el análisis de código,
            la lectura de métricas complementarias y la generación de reportes preliminares,
            manteniendo revisión humana antes de guardar conocimiento en memoria.
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown(
                card_html(
                    "▣",
                    "Insumos",
                    "Documento metodológico, código Python/SQL/notebook, observación del validador y métricas complementarias.",
                    "Entrada"
                ),
                unsafe_allow_html=True
            )

        with col2:
            st.markdown(
                card_html(
                    "◇",
                    "Preparación",
                    "Limpieza, anonimización básica y selección del contexto relevante.",
                    "Control"
                ),
                unsafe_allow_html=True
            )

        with col3:
            st.markdown(
                card_html(
                    "⌘",
                    "Tools",
                    "Funciones especializadas con BaseModel para revisar metodología, código, métricas y benchmark.",
                    "Ejecución"
                ),
                unsafe_allow_html=True
            )

        with col4:
            st.markdown(
                card_html(
                    "◉",
                    "Agente IA",
                    "Agente ReAct que razona sobre los insumos, usa tools y genera un reporte preliminar.",
                    "IA"
                ),
                unsafe_allow_html=True
            )

        st.markdown("<br>", unsafe_allow_html=True)

        col5, col6, col7, col8 = st.columns(4)

        with col5:
            st.markdown(
                card_html(
                    "▤",
                    "Evaluación",
                    "Clasifica hallazgos por severidad, impacto y necesidad de análisis adicional.",
                    "Riesgo"
                ),
                unsafe_allow_html=True
            )

        with col6:
            st.markdown(
                card_html(
                    "↻",
                    "Benchmark",
                    "Analiza escenarios alternativos cuando existe impacto potencial.",
                    "Comparación"
                ),
                unsafe_allow_html=True
            )

        with col7:
            st.markdown(
                card_html(
                    "☉",
                    "Revisión humana",
                    "El validador acepta, ajusta o descarta el hallazgo preliminar.",
                    "Control humano"
                ),
                unsafe_allow_html=True
            )

        with col8:
            st.markdown(
                card_html(
                    "▦",
                    "Memoria",
                    "Guarda feedback validado en SQLite para futuras revisiones.",
                    "Aprendizaje"
                ),
                unsafe_allow_html=True
            )

        st.markdown("""
        ### Componentes técnicos

        - `create_validation_agent()`: crea el agente IA.
        - `@tool`: define herramientas especializadas.
        - `BaseModel + Field`: estructura los inputs de las tools.
        - `SQLite`: memoria local del MVP.
        - `Streamlit`: interfaz de uso.
        - `LangSmith`: observabilidad avanzada opcional.
        - Trazabilidad local: visible en el panel izquierdo.
        - Soporte para múltiples archivos de código y `.zip`.
        """)

    with tab2:
        st.subheader("Carga de insumos")

        st.markdown("""
        <div class="mini-note">
        Formatos permitidos:<br>
        <b>Documento metodológico:</b> PDF, DOCX, TXT, MD.<br>
        <b>Código / notebook:</b> PY, IPYNB, SQL, TXT, MD, JSON o ZIP.<br>
        <b>Métricas complementarias:</b> CSV, XLSX, XLS.
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)

        with col1:
            archivo_metodologia = st.file_uploader(
                "Documento metodológico",
                type=FORMATOS_METODOLOGIA,
                key="metodologia"
            )

        with col2:
            archivos_codigo = st.file_uploader(
                "Código / Notebook / SQL / ZIP",
                type=FORMATOS_CODIGO,
                accept_multiple_files=True,
                key="codigo"
            )

        with col3:
            archivo_datos = st.file_uploader(
                "Métricas o dataset complementario",
                type=FORMATOS_DATOS,
                key="datos"
            )

        observacion_usuario = st.text_area(
            "Observación metodológica o foco de revisión",
            placeholder="Ejemplo: revisar si la población objetivo implementada en el código coincide con la metodología y si la caída del Gini podría representar degradación de performance.",
            height=120
        )

        st.markdown("""
        <div class="mini-note">
        Las métricas tabulares son complementarias. Si no se carga Excel/CSV, el estado de métricas quedará como Pendiente.
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button("Ejecutar revisión IA", type="primary"):
            try:
                limpiar_debug()

                st.session_state["evidencia_rag"] = ""

                log_debug("Inicio", f"Modelo={modelo}, Periodo={periodo}")

                actualizar_estado_flujo("insumos", "En proceso")
                actualizar_estado_flujo("preparacion", "Pendiente")
                actualizar_estado_flujo("metricas", "Pendiente")
                actualizar_estado_flujo("memoria", "Pendiente")
                actualizar_estado_flujo("agente", "Pendiente")
                actualizar_estado_flujo("reporte", "Pendiente")

                if not archivo_metodologia and not archivos_codigo and not observacion_usuario:
                    actualizar_estado_flujo("insumos", "Error")
                    log_debug(
                        "Validación de entrada",
                        "No se cargaron insumos ni observación.",
                        estado="Error"
                    )
                    st.error("Debes cargar al menos un documento/código o escribir una observación de validación.")
                    return

                texto_metodologia = ""
                codigo_modelo = ""
                resumen_metricas = ""

                with st.spinner("Leyendo y preparando insumos..."):

                    if archivo_metodologia:
                        texto_metodologia = leer_archivo(archivo_metodologia)
                        texto_metodologia = preparar_contexto(
                            texto_metodologia,
                            max_chars=MAX_CONTEXT_CHARS
                        )
                        log_debug(
                            "Documento leído",
                            f"Metodología preparada: {len(texto_metodologia)} caracteres"
                        )

                    if archivos_codigo:
                        codigo_modelo = leer_multiples_archivos_codigo(archivos_codigo)
                        codigo_modelo = preparar_contexto(
                            codigo_modelo,
                            max_chars=MAX_CONTEXT_CHARS
                        )
                        log_debug(
                            "Código leído",
                            f"Código preparado: {len(codigo_modelo)} caracteres desde {len(archivos_codigo)} archivo(s)"
                        )

                    if archivo_datos:
                        df = leer_excel_o_csv(archivo_datos)

                        if df is not None:
                            df_json = dataframe_to_json_records(df.head(500))
                            resumen_metricas = calcular_metricas_validacion.invoke({
                                "df_json": df_json
                            })

                            log_debug(
                                "Métricas calculadas",
                                "Tool calcular_metricas_validacion ejecutada"
                            )

                            actualizar_estado_flujo("metricas", "OK")
                            st.session_state["ultimo_resumen_metricas"] = resumen_metricas

                            resumen_gini = resumir_excel_gini(df)

                            if resumen_gini:
                                resumen_metricas = (
                                    resumen_metricas
                                    + "\n\nRESUMEN TENDENCIA GINI:\n"
                                    + resumen_gini
                                )
                                st.session_state["ultimo_resumen_metricas"] = resumen_metricas
                                log_debug(
                                    "Tendencia Gini",
                                    "Resumen de tendencia Gini generado"
                                )

                        else:
                            actualizar_estado_flujo("metricas", "Error")
                            log_debug(
                                "Métricas",
                                "No se pudo leer el archivo tabular.",
                                estado="Error"
                            )

                    else:
                        actualizar_estado_flujo("metricas", "Pendiente")
                        log_debug(
                            "Métricas",
                            "No se cargó archivo tabular complementario.",
                            estado="OK"
                        )

                    actualizar_estado_flujo("insumos", "OK")
                    actualizar_estado_flujo("preparacion", "OK")
                    log_debug("Preparación completada", "Contexto listo para el agente")

                texto_control = " ".join([

                    modelo or "",

                    periodo or "",

                    observacion_usuario or "",

                    texto_metodologia[:1000] or "",

                    codigo_modelo[:1000] or "",

                    resumen_metricas[:1000] or ""
                ])

                if not validar_dominio(texto_control):
                    actualizar_estado_flujo("reporte", "Error")
                    log_debug(
                        "Dominio",
                        "Consulta fuera del alcance del MVP.",
                        estado="Error"
                    )
                    st.warning(
                        "La consulta parece estar fuera del dominio del MVP. "
                        "Este asistente está diseñado para validación de modelos de riesgo en banca."
                    )
                    return

                with st.spinner("Consultando memoria histórica..."):
                    keyword_memoria = observacion_usuario if observacion_usuario else modelo
                    memoria_historica = consultar_hallazgos_previos(
                        keyword=keyword_memoria,
                        limite=5
                    )

                    actualizar_estado_flujo("memoria", "OK")
                    log_debug(

                        "Memoria consultada",

                        "SQLite respondió correctamente"

                    )

                    pregunta_rag = (

                         observacion_usuario

                         if observacion_usuario

                         else modelo

                    )

                    contexto_rag = obtener_contexto_rag(

                         pregunta_rag

                    )

                    st.session_state["evidencia_rag"] = contexto_rag

                    actualizar_estado_flujo(

                        "rag",

                         "OK"

                    )

                    actualizar_estado_flujo(

                        "evidencia",

                        "OK"

                    )

                with st.expander("Vista previa de contexto preparado"):
                    st.write("Documento metodológico:")
                    st.text(
                        texto_metodologia[:2000]
                        if texto_metodologia
                        else "No cargado."
                    )

                    st.write("Código:")
                    st.text(
                        codigo_modelo[:1500]
                        if codigo_modelo
                        else "No cargado."
                    )

                    st.write("Métricas:")
                    st.text(
                        resumen_metricas
                        if resumen_metricas
                        else "No se cargaron métricas complementarias."
                    )

                    st.write("Memoria:")
                    st.text(memoria_historica)

                with st.spinner("Creando agente y generando reporte preliminar..."):
                    agent = create_validation_agent()

                    actualizar_estado_flujo("agente", "OK")
                    log_debug("Agente creado", f"Backend: {AGENT_BACKEND}")

                    prompt = generar_prompt_validacion(
                        modelo=modelo,
                        periodo=periodo,
                        texto_metodologia=texto_metodologia,
                        codigo_modelo=codigo_modelo,
                        resumen_metricas=resumen_metricas,
                        memoria_historica=memoria_historica,
                        observacion_usuario=observacion_usuario
                    )

                    st.session_state["ultimo_prompt"] = prompt

                    log_debug(
                        "Prompt generado",
                        f"{len(prompt)} caracteres enviados al agente"
                    )

                    reporte = ejecutar_agente(agent, prompt)

                    actualizar_estado_flujo("reporte", "OK")
                    log_debug(
                        "Reporte generado",
                        f"{len(reporte)} caracteres recibidos"
                    )

                st.session_state["reporte_preliminar"] = reporte
                st.session_state["modelo"] = modelo
                st.session_state["periodo"] = periodo
                st.session_state["observacion_usuario"] = observacion_usuario

##### se borro postgresql

                st.success("Reporte preliminar generado.")
                st.markdown("## Reporte preliminar")
                st.markdown(reporte)

            except Exception as e:
                error_msg = traceback.format_exc()

                st.session_state["ultimo_error"] = error_msg

                log_debug("Error", str(e), estado="Error")
                actualizar_estado_flujo("reporte", "Error")

                guardar_ejecucion(
                    modelo=modelo,
                    periodo=periodo,
                    estado="ERROR",
                    resumen="Error durante ejecución del análisis.",
                    error=error_msg
                )

                st.error(f"Ocurrió un error: {str(e)}")

                with st.expander("Detalle técnico"):
                    st.code(error_msg)

    with tab3:
        st.subheader("Validación humana y guardado en memoria")

        reporte = st.session_state.get("reporte_preliminar", "")

        if not reporte:
            st.warning("Primero ejecuta una revisión en la pestaña Revisión IA.")
        else:
            st.markdown("### Reporte preliminar generado")
            st.markdown(reporte)

            st.divider()

            st.markdown("### Registrar feedback validado")

            categoria = st.selectbox(
                "Categoría del hallazgo",
                [
                    "Metodología",
                    "Código / implementación",
                    "Población objetivo",
                    "Métricas / resultados",
                    "Trazabilidad",
                    "Benchmark",
                    "Documentación",
                    "Otro"
                ]
            )

            hallazgo = st.text_area(
                "Hallazgo validado por el humano",
                placeholder="Escribe el hallazgo final que sí deseas guardar en memoria.",
                height=120
            )

            col_a, col_b = st.columns(2)

            with col_a:
                severidad = st.selectbox(
                    "Severidad",
                    ["Baja", "Media", "Alta"]
                )

            with col_b:
                impacto = st.selectbox(
                    "Impacto",
                    [
                        "Sin impacto material",
                        "Impacto documental",
                        "Impacto metodológico",
                        "Impacto en resultados",
                        "Impacto por confirmar"
                    ]
                )

            recomendacion = st.text_area(
                "Recomendación",
                placeholder="Acción sugerida para el equipo de desarrollo/modelamiento.",
                height=100
            )

            decision_humana = st.selectbox(
                "Decisión humana",
                ["Aceptar", "Ajustar", "Descartar"]
            )

            comentario_humano = st.text_area(
                "Comentario del validador",
                placeholder="Comentario adicional, sustento o aclaración.",
                height=100
            )

            confirmar = st.checkbox(
                "Confirmo que este feedback fue revisado por un validador humano y puede guardarse en memoria."
            )

            if st.button("Guardar feedback validado", type="primary"):
                if not confirmar:
                    st.error("Debes confirmar la revisión humana antes de guardar.")

                elif not hallazgo.strip():
                    st.error("Debes ingresar un hallazgo validado.")

                else:
                    guardar_hallazgo_validado(
                        modelo=st.session_state.get("modelo", modelo),
                        periodo=st.session_state.get("periodo", periodo),
                        categoria=categoria,
                        hallazgo=hallazgo,
                        severidad=severidad,
                        impacto=impacto,
                        recomendacion=recomendacion,
                        decision_humana=decision_humana,
                        comentario_humano=comentario_humano
                    )

                    log_debug(
                        "Memoria actualizada",
                        "Feedback humano guardado en SQLite"
                    )

                    actualizar_estado_flujo(

                        "memoria",

                        "OK"

                    )

                    st.success("Feedback validado guardado en SQLite.")

    with tab4:
        st.subheader("Memoria del asistente")

        tabla = st.selectbox(
            "Selecciona tabla",
            ["hallazgos", "ejecuciones"]
        )

        if st.button("Actualizar tabla"):
            try:
                df_mem = obtener_tabla(tabla)
                st.dataframe(df_mem, use_container_width=True)

            except Exception as e:
                st.error(f"No se pudo leer la tabla: {str(e)}")

        st.caption("Si el texto se ve truncado, puedes descargar la tabla como CSV.")

        try:
            df_mem_download = obtener_tabla(tabla)
            csv = df_mem_download.to_csv(index=False).encode("utf-8")

            st.download_button(
                label="Descargar tabla como CSV",
                data=csv,
                file_name=f"{tabla}.csv",
                mime="text/csv"
            )

        except Exception:
            pass


if __name__ == "__main__":
    main()
